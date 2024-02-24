import os
import re
import docx
import pathlib

from docx.shared import Cm, Inches, RGBColor, Pt


def find_number_of_pupils(doc_file):
    no_of_pupils = 0
    
    doc = docx.Document(doc_file)
    paragraphs = doc.paragraphs

    name_patt = re.compile("NAME:\s*([\w\s\W]+)")

    for paragraph in paragraphs:
        pupil_name = name_patt.match(paragraph.text)
                                    
        if pupil_name:
            no_of_pupils += 1

        else:
            continue

    return no_of_pupils


def search_school_files(path):
    files = []
    non_files = []
    
    pattern1 = re.compile("([\d]{8})\s+([A-Z\s\W]+)\s*pg\s*[\d]\s*-\s*[\d]+\s*\.docx")
    pattern2 = re.compile("([\d]{8})\s+([A-Z\s\W]+)\s*[\d]\s*-\s*[\d]+\s*\.docx")
    pattern3 = re.compile("([\d]{8})\s+([A-Z\s\W]+)\s*\.docx")
    
    for _file in os.listdir(path):
        if pattern1.match(_file):
            file_data = pattern1.match(_file)
            code, school = file_data.groups()
            no_of_pupils = find_number_of_pupils(os.path.join(path, _file))
            
            files.append((int(code), str(school), no_of_pupils, "", "", ""))

        else:
            if pattern2.match(_file):
                file_data = pattern2.match(_file)
                code, school = file_data.groups()
                no_of_pupils = find_number_of_pupils(os.path.join(path, _file))
            
                files.append((int(code), str(school), no_of_pupils, "", "", ""))

            else:
                if pattern3.match(_file):
                    file_data = pattern3.match(_file)
                    code, school = file_data.groups()
                    no_of_pupils = find_number_of_pupils(os.path.join(path, _file))
            
                    files.append((int(code), str(school), no_of_pupils, "", "", ""))

                else:
                    non_files.append(_file)
                
                    continue

    return files, non_files


def write_to_word(path, records, county, subcounty):
    doc = docx.Document()

    doc.add_heading(f"SCHOOLS CHECKLIST", 0)
    
    paragraph = doc.add_paragraph()
    paragraph2 = paragraph.add_run(f"{county.upper()} COUNTY, {subcounty.upper()} SUB-COUNTY")
    paragraph2.underline = True
    paragraph2.font.size = Pt(16)
    
    table = doc.add_table(rows = 1, cols = 7)
    table.style = doc.styles["Light Grid Accent 1"]

    table.columns[0].width = Cm(1.0)
    table.columns[1].width = Cm(1.9)
    table.columns[2].width = Cm(8.5)
    table.columns[3].width = Cm(1.5)
    table.columns[4].width = Cm(1.5)
    table.columns[5].width = Cm(2.0)
    table.columns[6].width = Cm(1.5)

    cells = table.rows[0].cells
    
    cells[0].text = "NO."
    cells[0].width = Cm(1.5)
    
    cells[1].text = "CODE"
    cells[1].width = Cm(1.9)
    
    cells[2].text = "SCHOOL NAME"
    cells[2].width = Cm(8.5)
    
    cells[3].text = "PUPILS"
    cells[3].width = Cm(1.5)
    
    cells[4].text = "CHECK"
    cells[4].width = Cm(1.5)

    cells[5].text = "RECIPIENT"
    cells[5].width = Cm(2.0)

    cells[6].text = "SIGN."
    cells[6].width = Cm(1.5)

    for i, row in enumerate(records, 1):
        code, school, num, rec, check, sign = row
        
        row_cells = table.add_row().cells
        
        row_cells[0].text = str(i)
        row_cells[1].text = str(code)
        row_cells[2].text = school
        row_cells[3].text = str(num)
        row_cells[4].text = check
        row_cells[5].text = rec
        row_cells[6].text = sign

    doc_file = os.path.join(path, f"{county.upper()}_{subcounty.upper()}_CHECKLIST.docx")

    doc.save(doc_file)
    

while True:
    print(" ------\n")
    path_input = str(input(" [*] Enter School Files Directory: "))
    path_input = pathlib.Path(path_input)

    if not os.path.exists(path_input):
        print(" [!] Incorrect School Directory!!")
        print("")
        
        continue
    
    county = str(input(" [*] Enter County (e.g. Nairobi): "))
    subcounty = str(input(" [*] Enter SubCounty (e.g. Kibra): "))

    records, nfiles = search_school_files(path_input)

    if nfiles:
        print(f" [+] Exempted files ({len(nfiles)}):")
        print(" ---")
        
        for file in nfiles:
            print(f"    - {file}")

        print("")
        
    write_to_word(path_input, records, county, subcounty)

    print(" [+] Check list generated successfully.")




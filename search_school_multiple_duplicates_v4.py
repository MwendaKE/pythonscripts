import os
import docx
import re
import pathlib

from collections import deque


def search_doc_files(path):
    doc_files = []

    for file in os.listdir(path):
        if re.match("\d{8}\s+[\w\W\d\s]+", file) and file.endswith(".docx"):
            doc_files.append(os.path.join(path, file))

        else:
            continue

    return doc_files


def search_stud_names(doc_file):
    stud_names = []

    doc = docx.Document(doc_file)
    paragraphs = doc.paragraphs

    name_patt = re.compile("NAME:\s*([\w\s\W]+)")

    for paragraph in paragraphs:
        stud_name = name_patt.match(paragraph.text)
    
        if stud_name:
            stud_names.append(stud_name.groups()[0])

        else:
            continue

    return stud_names


def search_schoolcodes_in_content(doc_file):
    '''this code ensures that the filename matches the name of school
        in the file content'''
    school_codes = deque()

    try:
        doc = docx.Document(doc_file)
        
    except Exception as e:
        import sys
        
        print(f" [+] Error: {e}")
        sys.exit()

    paragraphs = doc.paragraphs
    
    sch_pattern = re.compile("SCHOOL:\s*([\d]{8})\s*([\w\s\W]+)")
        
    for paragraph in paragraphs:
        school_name = sch_pattern.match(paragraph.text)

        if school_name:
            school_code = school_name.groups()[0]

            school_codes.append(school_code)

        else:
            continue

    return school_codes


def search_tel_no_present(doc_file):
    doc = docx.Document(doc_file)
    paragraphs = doc.paragraphs

    tel_patt = re.compile("(TEL:\s*[\d]{10}\s*/\s*[\d]{10}\.)")

    for paragraph in paragraphs:
        tel_no = tel_patt.match(paragraph.text)

        if not tel_no:
            print(" ! Tel No. not found!")
            #print(f" -> {tel_no.groups()[0]}")
            


def search_duplicate_names(stud_names):
    duplicates = []
    
    for name in  stud_names:
        if stud_names.count(name) >= 2:
            duplicates.append(name)

        else:
            continue

    return duplicates


while True:
    print("-" * 20)
    schs_path = str(input(" [+] Enter Schools Directory: "))

    if not os.path.exists(pathlib.Path(schs_path)):
        print(" [!] Incorrect School Directory!!")
        continue

    
    pattern1 = re.compile("([\d]{8})\s*[A-Z\s\W]+\s*pg\s*[\d]\s*-\s*[\d]+")
    pattern2 = re.compile("([\d]{8})\s*[A-Z\s\W]+\s*[\d]\s*-\s*[\d]+")

    duplicates_file = os.path.join(schs_path, "Duplicate Analysis.txt")
    
    try:
        prev_dir = os.path.split(os.path.split(duplicates_file)[0])[1]
        duplicates_file = os.path.join(schs_path, f"DUPLICATE ANALYSIS - {prev_dir.upper()}.txt")

    except :
        pass
        

    doc_files = search_doc_files(schs_path)
    
    with open(duplicates_file, "w") as txtfile:
        print("\n" + "-" * 20)
        print("\n [+] Writing to text file...")
        print("\n" + "-" * 20)

        for doc_file in doc_files:
            stud_names = search_stud_names(doc_file)
            student_duplicates = search_duplicate_names(stud_names)

            search_tel_no_present(doc_file)

            file_len = len(os.path.basename(doc_file))
        
            print(f" [*] Processing '{os.path.basename(doc_file).upper()}' ...")
         
            txtfile.write("\n" + "=" * (file_len + 5) + "\n")
            txtfile.write(f" |{os.path.basename(doc_file).upper()} |")
            txtfile.write("\n" + "=" * (file_len + 5) + "\n")
            txtfile.write(f" Student Number : {len(stud_names)}\n")
            txtfile.write(f" Student Name Duplicates: {len(student_duplicates)}\n")
        
            if student_duplicates:
                txtfile.write("-" * 20 + "\n")
                
                for duplicate in student_duplicates:
                    txtfile.write(f" - {duplicate}\n")

                txtfile.write("\n")

            base_doc_file = os.path.basename(doc_file)
            
            if pattern1.match(base_doc_file):
                file_sch_code = pattern1.match(base_doc_file).groups()[0]

            else:
                if pattern2.match(doc_file):
                    file_sch_code = pattern2.match(base_doc_file).groups()[0]

                else:
                    continue

            school_codes = search_schoolcodes_in_content(doc_file)

            txtfile.write(" School Code Duplicates:\n")
            txtfile.write("-" * 20 + "\n")
            
            for i, content_sch_code in enumerate(school_codes, 1):
                if file_sch_code != content_sch_code:
                    txtfile.write(f" - PG {i}: {file_sch_code} -> {content_sch_code} \n")

                else:
                    continue


                
    print("")
    print(" [+] Code Analysis successful.")
    print("")
